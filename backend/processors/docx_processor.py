import docx
from datetime import datetime

def process_docx(file_path):
    doc = docx.Document(file_path)
    
    # Extract paragraphs and headings
    paragraphs = [p.text for p in doc.paragraphs]
    
    # Get a preview of the content
    preview_text = ' '.join(paragraphs[:3])[:500] + "..." if paragraphs else ""
    
    # Count paragraphs, tables, etc.
    table_count = len(doc.tables)
    paragraph_count = len(paragraphs)
    
    return {
        'type': 'docx',
        'preview': preview_text,
        'paragraph_count': paragraph_count,
        'table_count': table_count,
        'processed_at': datetime.now().isoformat()
    } 