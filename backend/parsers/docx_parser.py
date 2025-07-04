import os
import logging
from typing import Dict, List, Any
from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """Extract full text from a DOCX file, cleaning up unnecessary whitespace."""
    try:
        doc = Document(file_path)
        return "\n".join(filter(None, (para.text.strip() for para in doc.paragraphs)))

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_tables_with_context(doc: Document) -> List[Dict[str, Any]]:
    """Extract tables with preceding paragraph as title where possible."""
    tables_data = []
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    for i, table in enumerate(doc.tables):
        prev_title = paragraphs[i - 1] if i > 0 and i - 1 < len(paragraphs) else f"Table {i + 1}"
        table_content = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables_data.append({"title": prev_title, "content": table_content})
    return tables_data

def extract_metadata_from_docx(file_path: str) -> Dict[str, Any]:
    """Extract metadata including author, file size, and document statistics."""

    try:
        doc = Document(file_path)
        metadata = {
            "filename": os.path.basename(file_path),
            "filesize": os.path.getsize(file_path),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "word_count": sum(len(p.text.split()) for p in doc.paragraphs),
            "page_count": len(doc.sections),
        }
        try:
            core_props = doc.core_properties
            metadata.update({
                "author": core_props.author or "Not available",
                "title": core_props.title or "Not available",
                "created": core_props.created.isoformat() if core_props.created else "Not available",
                "modified": core_props.modified.isoformat() if core_props.modified else "Not available",
                "last_modified_by": core_props.last_modified_by or "Not available",
                "revision": core_props.revision or "Not available"
            })
        except Exception as e:
            logger.warning(f"Error extracting core properties: {str(e)}")
        return metadata
    except Exception as e:
        logger.error(f"Error extracting metadata from DOCX: {str(e)}")
        return {}

def extract_sections_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """Extract structured sections from a DOCX document."""
    doc = Document(file_path)
    sections, current_section, current_content = [], None, []
    common_headings = {"Personas", "Gobierno", "Seguridad", "Cultura", "Medios de vida", "Infraestructura", "Medio ambiente", "Tierra y recursos naturales"}
    invalid_headings = {"N/C", "N/S", "Completar con:"}  # exclude these headings
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        is_heading = (para.style and 'Heading' in para.style.name) or text.isupper() or text.endswith(":") or text in common_headings
        if is_heading and text not in invalid_headings:
            if current_section and current_content:
                sections.append({"heading": current_section, "content": "\n".join(current_content)})
            current_section, current_content = text, []
        else:
            current_content.append(text)
    
    if current_section and current_content:
        sections.append({"heading": current_section, "content": "\n".join(current_content)})
    return sections

def parse_docx(file_path: str) -> Dict[str, Any]:
    """Parse DOCX file into structured metadata, text, sections, and tables."""
    doc = Document(file_path)
    return {
        "metadata": extract_metadata_from_docx(file_path),
        "text": extract_text_from_docx(file_path),
        "sections": extract_sections_from_docx(file_path),
        "tables": extract_tables_with_context(doc)
    }
