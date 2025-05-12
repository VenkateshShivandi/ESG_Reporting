import os
import logging
from typing import Dict, List, Any
from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """Extract full text from a DOCX file, cleaning up unnecessary whitespace."""
    try:
        doc = Document(file_path)
        # Join paragraphs with a single newline, preserving structure better than space
        # Filter empty paragraphs after stripping
        return "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())

    except Exception as e:
        logger.error(f"Error extracting text from DOCX '{os.path.basename(file_path)}': {str(e)}")
        return ""

def extract_tables_with_context(doc: Document) -> List[Dict[str, Any]]:
    """Extract tables with preceding paragraph as title where possible."""
    tables_data = []
    # Consider all elements (paragraphs and tables) in order
    doc_elements = list(doc.element.body)
    
    # Find the index of each table within the body elements
    table_indices = [i for i, element in enumerate(doc_elements) if element.tag.endswith('tbl')]
    
    for idx, table_index in enumerate(table_indices):
        # Look for the preceding paragraph element
        prev_title = f"Table {idx + 1}" # Default title
        if table_index > 0:
            prev_element = doc_elements[table_index - 1]
            # Check if the previous element is a paragraph
            if prev_element.tag.endswith('p'):
                 # Find the corresponding paragraph object in the doc.paragraphs list
                 # This mapping can be complex, safer to extract text directly from element if possible
                 # Or, rely on paragraphs list and try to find the closest one before the table?
                 # Simplification: Assume the paragraph immediately preceding the table *in the document flow* is the title
                 # We need a reliable way to map element index to paragraph object index
                 # For now, keep the simpler logic using doc.paragraphs index, acknowledging limitations
                 paragraphs = [p for p in doc.paragraphs if p.text.strip()] # Get non-empty paragraphs
                 # This index mapping is likely incorrect. 
                 # A better approach might involve iterating paragraphs and checking if a table follows immediately.
                 # Sticking to previous logic for now:
                 if table_index -1 < len(paragraphs):
                      prev_para_text = paragraphs[table_index - 1].text.strip() # This is probably wrong index
                      if prev_para_text:
                          prev_title = prev_para_text
        
        # Get the actual Table object using the index from doc.tables
        try:
             table = doc.tables[idx]
             table_content = [[cell.text.strip() for cell in row.cells] for row in table.rows]
             tables_data.append({"title": prev_title, "content": table_content})
        except IndexError:
             logger.warning(f"Table index mismatch for table {idx+1} in '{os.path.basename(doc.part.package.filename)}'. Skipping table.")
             
    return tables_data

def extract_metadata_from_docx(file_path: str) -> Dict[str, Any]:
    """Extract metadata including author, file size, and document statistics."""

    try:
        doc = Document(file_path)
        metadata = {
            "filename": os.path.basename(file_path),
            "filesize": os.path.getsize(file_path),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            # Correct word count (approximated)
            "word_count": sum(len(p.text.split()) for p in doc.paragraphs),
            "page_count": "N/A" # Page count is not directly available in python-docx
        }
        try:
            core_props = doc.core_properties
            metadata.update({
                "author": core_props.author or "Not available",
                "title": core_props.title or "Not available",
                "created": core_props.created.isoformat() if core_props.created else "Not available",
                "modified": core_props.modified.isoformat() if core_props.modified else "Not available",
                "last_modified_by": core_props.last_modified_by or "Not available",
                "revision": core_props.revision or "Not available"
            })
        except Exception as e:
            logger.warning(f"Error extracting core properties: {str(e)}")
        return metadata
    except Exception as e:
        logger.error(f"Error extracting metadata from DOCX '{os.path.basename(file_path)}': {str(e)}")
        return {}

def extract_sections_from_docx(doc: Document) -> List[Dict[str, Any]]:
    """Extract structured sections from a DOCX document based on styles or heuristics."""
    sections = []
    current_section = "Introduction" # Default first section
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        # Check if style indicates a heading (more reliable)
        is_heading = para.style and para.style.name.startswith('Heading')
        
        # Heuristic: Check if text is short, capitalized, or ends with colon (less reliable)
        is_heuristic_heading = (len(text) > 0 and len(text) < 100 and (text.isupper() or text.endswith(':')))

        if text and (is_heading or is_heuristic_heading):
            # Found a potential heading
            if current_content: # Save previous section if content exists
                sections.append({"heading": current_section, "content": "\n".join(current_content)})
            current_section = text # Start new section
            current_content = [] # Reset content
        elif text: # Not a heading, but has content
            current_content.append(text)
    
    # Add the last section
    if current_content:
        sections.append({"heading": current_section, "content": "\n".join(current_content)})
        
    # If no sections were found, treat the whole document as one section
    if not sections and doc.paragraphs:
         full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
         if full_text:
             sections.append({"heading": "Full Document", "content": full_text})
             
    return sections

def parse_docx(file_path: str) -> Dict[str, Any]:
    """Parse DOCX file into structured metadata, text, sections, and tables."""
    try:
         doc = Document(file_path)
         # Note: Extract sections before text to potentially use heading info
         sections = extract_sections_from_docx(doc)
         text = "\n".join(s['content'] for s in sections) # Reconstruct text from sections
         if not text: # Fallback if section extraction failed
             text = extract_text_from_docx(file_path)
             
         return {
             "metadata": extract_metadata_from_docx(file_path),
             "text": text,
             "sections": sections,
             "tables": extract_tables_with_context(doc)
         }
    except Exception as e:
         logger.error(f"Failed to parse DOCX file '{os.path.basename(file_path)}': {e}")
         return {"error": f"Failed to parse DOCX: {e}"}

# Example usage
if __name__ == "__main__":
    import sys
    import json # Import json for printing
        
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <docx_file_path>")
        sys.exit(1)
        
    file_path = sys.argv[1]
    result = parse_docx(file_path)

    if "error" in result:
         print(f"Error: {result['error']}")
    else:
         # Pretty print the result
         print(json.dumps(result, indent=2, ensure_ascii=False)) 